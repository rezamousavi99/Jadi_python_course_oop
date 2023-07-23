from docx import Document

class DocxNumberInverter:
    def __init__(self, path) -> None:
        self.document = Document(path)
        self.numbersdone = 0
        self.numberoffixes = 0
        self.newtext = ''
    
    def invert_numbers(self):
        for para in self.document.paragraphs:
            self.newtext = ''
            for char in para.text:
                if char.isnumeric():
                    self.numberoffixes += 1
                    if self.numbersdone == 0:
                        self.newtext += char
                        self.numbersdone += 1
                    else:
                        self.newtext = self.newtext[:-1*self.numbersdone] + char + self.newtext[-1*self.numbersdone:]
                        self.numbersdone += 1
                
                else:
                    self.newtext += char
                    self.numbersdone = 0
            
            para.text = self.newtext
        print('done...')
            # print(para.text)
    def display_text(self):
        for para in self.document.paragraphs:
            print(para.text)

    def save_file(self, location):
        self.document.save(location)
        print(f'document saved at {location} successfully...')

def commands():
    print('1.invert the numbers in the text\n2.display content of document\n3.save the document\n4.exit')

def main():
    uinput = input('Enter docx file path: ')
    d = DocxNumberInverter(uinput)

    while True:
        commands()
        inpt = input('enter a command: ')
        match inpt:
            case '1':
                d.invert_numbers()

            case '2':
                d.display_text()
            
            case '3':
                d.save_file('saveddeom.docx')
                break
            
            case '4':
                print('exiting...')
                break

if __name__ == '__main__':
    main()
    